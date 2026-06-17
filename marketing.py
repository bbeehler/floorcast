# marketing.py
import streamlit as st
import pandas as pd
import datetime
import plotly.express as px
import plotly.graph_objects as go
import google.generativeai as genai
import uuid
import time

# Try to import python-docx for the bulk loader
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- SHARED HELPER FUNCTIONS ---

def archive_sentiment_entry(text, asset_tag, review_date, tenant_id, supabase):
    """Passes review text to Gemini to generate a sentiment score and vaults it."""
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"Analyze the sentiment of this guest review. Return ONLY a single float between -1.0 (extremely negative) and 1.0 (extremely positive). Review: {text}"
        res = model.generate_content(prompt)
        
        try:
            score = float(res.text.strip())
        except ValueError:
            score = 0.0

        if score > 0.3: cat = "Positive"
        elif score < -0.3: cat = "Negative"
        else: cat = "Neutral"

        payload = {
            "id": str(uuid.uuid4()),
            "tenant_id": tenant_id,
            "asset": asset_tag,
            "sentiment_score": score,
            "sentiment_category": cat,
            "raw_text": text,
            "timestamp": review_date.strftime("%Y-%m-%dT12:00:00")
        }
        supabase.table("mt_sentiment_history").insert(payload).execute()
        return True
    except Exception as e:
        st.error(f"Archival Sync Error: {e}")
        return False

def get_forensic_metrics(df_input, coeffs):
    if not df_input: return {"df": pd.DataFrame()}
    df = pd.DataFrame(df_input).copy()
    df['entry_date'] = pd.to_datetime(df['entry_date'])
    
    # 1. Heartbeat Baseline
    hb = {d: float(coeffs.get(f'{d[:3]}_Base', 5000)) for d in ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']}
    df['baseline'] = df['entry_date'].dt.day_name().map(hb).astype(float)
    
    # 2. Marketing Attribution Logic
    dec = float(coeffs.get('Ad_Decay', 85))/100
    c1 = float(coeffs.get('Clicks', 0.05))
    c2 = float(coeffs.get('Social_Imp', 0.0002))
    
    pool, lift = 0.0, []
    for _, r in df.iterrows():
        pool = ((float(r.get('ad_clicks', 0) or 0)*c1) + (float(r.get('ad_impressions', 0) or 0)*c2)) + (pool * dec)
        lift.append(pool)
    df['residual_lift'] = lift
    
    # 3. Final Synthesis
    if 'predicted_traffic' in df.columns:
        fallback_calc = df['baseline'] + df['residual_lift'] + float(coeffs.get('Promo', 500.0))
        df['expected'] = pd.to_numeric(df['predicted_traffic'], errors='coerce').fillna(fallback_calc)
    else:
        df['expected'] = df['baseline'] + df['residual_lift'] + float(coeffs.get('Promo', 500.0))
        
    return {"df": df}


# --- MAIN MODULE RENDER ---
def render_marketing_module(supabase, tenant_id, property_name):
    st.title("📈 Marketing & Brand Intelligence")
    st.write(f"Advanced demand generation and sentiment analytics for **{property_name}**.")

    tabs = st.tabs(["🎯 Multi-Touch Attribution", "🔮 Scenario Simulator", "🧠 Brand Sentiment Vault"])

    # --- FETCH SECURE TENANT DATA ---
    res = supabase.table("mt_ledger").select("*").eq("tenant_id", tenant_id).order("entry_date", desc=True).execute()
    ledger_data = res.data if res.data else []
    
    # In a full SaaS, coefficients should also be a table tied to tenant_id. 
    # For now, we use a safe default dictionary to keep the math running.
    current_weights = {
        'Ad_Decay': 85, 'Clicks': 0.05, 'Social_Imp': 0.0002, 
        'Broadcast_Weight': 150, 'OOH_Weight': 100, 'Avg_Coin_In': 112.50
    }

    # ==========================================
    # TAB 1: ATTRIBUTION ANALYTICS
    # ==========================================
    with tabs[0]:
        if not ledger_data:
            st.info("💡 Forensic Vault empty. Populate the Casino Ledger to unlock attribution modeling.")
        else:
            m_full = get_forensic_metrics(ledger_data, current_weights)
            df_full_attr = m_full.get('df', pd.DataFrame())
            
            if not df_full_attr.empty:
                # Dynamic Date Filter
                df_full_attr['entry_date'] = pd.to_datetime(df_full_attr['entry_date'])
                min_date = df_full_attr['entry_date'].min().date()
                max_date = df_full_attr['entry_date'].max().date()

                col_d, _ = st.columns([1.5, 2.5])
                with col_d:
                    audit_window = st.date_input("Select Attribution Window:", value=(min_date, max_date), key="mktg_attr_window")

                if isinstance(audit_window, tuple) and len(audit_window) == 2:
                    start_date, end_date = audit_window
                    mask = (df_full_attr['entry_date'].dt.date >= start_date) & (df_full_attr['entry_date'].dt.date <= end_date)
                    df_attr = df_full_attr.loc[mask].copy()
                else:
                    df_attr = df_full_attr.copy()

                if df_attr.empty:
                    st.warning("No data falls within the selected window.")
                else:
                    # Data Structure Insurance
                    required_columns = {'actual_traffic': 0.0, 'baseline': 5000.0, 'residual_lift': 0.0, 'gravity_lift': 0.0, 'ad_clicks': 0.0}
                    for col, default_val in required_columns.items():
                        if col not in df_attr.columns: df_attr[col] = default_val
                        else: df_attr[col] = pd.to_numeric(df_attr[col], errors='coerce').fillna(default_val)
                    
                    # Component Calcs
                    total_guests = df_attr['actual_traffic'].sum()
                    organic_base = df_attr['baseline'].sum()
                    digital_lift = df_attr['residual_lift'].sum()
                    gravity_lift = df_attr['gravity_lift'].sum()
                    brand_inertia = (current_weights.get('Broadcast_Weight', 150) + current_weights.get('OOH_Weight', 100)) * len(df_attr)

                    st.markdown("### 🕰️ Multi-Touch Attribution (Time Decay Model)")
                    mta_cols = st.columns(3)
                    mta_cols[0].metric("Last-Touch (Digital)", f"{digital_lift:,.0f}")
                    mta_cols[1].metric("Assisted (Brand)", f"{brand_inertia:,.0f}")
                    mta_cols[2].metric("Conversion (Gravity)", f"{gravity_lift:,.0f}")
                    
                    st.divider()
                    
                    st.markdown("### 📡 Offline-to-Online Attribution Contribution")
                    col_pie, col_water = st.columns([1, 1.5])
                    with col_pie:
                        pie_labels = ['Organic (Baseline)', 'Online (Digital)', 'Offline (Brand)', 'Event Gravity']
                        pie_values = [organic_base, digital_lift, brand_inertia, gravity_lift]
                        fig_pie = px.pie(names=pie_labels, values=pie_values, hole=0.6)
                        fig_pie.update_layout(showlegend=True, height=350, margin=dict(l=0, r=0, t=0, b=0))
                        st.plotly_chart(fig_pie, use_container_width=True)

                    with col_water:
                        fig_water = go.Figure(go.Waterfall(
                            orientation="v",
                            measure=["relative", "relative", "relative", "relative", "total"],
                            x=["Organic", "Offline Media", "Online Signal", "Event Gravity", "Total Floor"],
                            y=[organic_base, brand_inertia, digital_lift, gravity_lift, total_guests]
                        ))
                        fig_water.update_layout(height=350, margin=dict(l=10, r=10, t=10, b=10))
                        st.plotly_chart(fig_water, use_container_width=True)

    # ==========================================
    # TAB 2: SCENARIO SIMULATOR
    # ==========================================
    with tabs[1]:
        st.subheader("Predictive Scenario Simulator")
        with st.container(border=True):
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                sim_date = st.date_input("Target Date", value=datetime.date.today() + datetime.timedelta(days=14))
                sim_season = st.selectbox("Business Season", ["Winter", "Spring", "Summer", "Autumn", "Peak"])
            with c2:
                sim_event = st.number_input("Event Attendance", value=0, step=500)
                sim_clicks = st.number_input("Planned Ad Clicks", value=1000, step=100)
            with c3:
                sim_imps = st.number_input("Planned Impressions", value=50000, step=5000)
                sim_rain = st.slider("Rain (mm)", 0, 50, 0)
            with c4:
                sim_snow = st.slider("Snow (cm)", 0, 30, 0)

            run_sim = st.button("🚀 Run Seasonal Scenario Projection", use_container_width=True)

        if run_sim:
            try:
                dow = sim_date.strftime('%A')
                
                # Baseline Logic
                if 'df_full_attr' in locals() and not df_full_attr.empty:
                    dow_history = df_full_attr[df_full_attr['entry_date'].dt.day_name() == dow].copy()
                    lifetime_baseline = dow_history['actual_traffic'].mean() if not dow_history.empty else 1500
                else:
                    lifetime_baseline = 1500

                # Multipliers
                seasonal_map = {"Winter": 0.85, "Spring": 1.05, "Summer": 1.15, "Autumn": 1.20, "Peak": 1.35}
                season_mult = seasonal_map.get(sim_season, 1.0)
                seasonal_baseline = lifetime_baseline * season_mult

                digital_lift = (sim_clicks * current_weights.get('Clicks', 0.05)) + (sim_imps * current_weights.get('Social_Imp', 0.0002))
                gravity_lift = sim_event * 0.25
                friction = (sim_rain * -12) + (sim_snow * -45)

                projected_guests = max(0, seasonal_baseline + digital_lift + gravity_lift + friction)
                projected_rev = projected_guests * current_weights.get('Avg_Coin_In', 112.50)
                
                st.divider()
                res1, res2, res3, res4 = st.columns(4)
                res1.metric(f"Lifetime {dow}", f"{lifetime_baseline:,.0f}")
                res2.metric("Seasonal Base", f"{seasonal_baseline:,.0f}")
                res3.metric("AI Projection", f"{projected_guests:,.0f}")
                res4.metric("Proj. Revenue", f"${projected_rev:,.0f}")

            except Exception as e:
                st.error(f"Simulation Error: {e}")

    # ==========================================
    # TAB 3: BRAND SENTIMENT VAULT (NEW)
    # ==========================================
    with tabs[2]:
        st.markdown("<h3 style='color: #111827; margin-top: 1rem;'>Vault Research & Archival</h3>", unsafe_allow_html=True)
        
        # Hardcoded standard assets (can link to DB later)
        tags = ["Overall Property", "Casino Floor", "Hotel / Lodging", "Food & Beverage", "Staff / Service"]

        # --- 1. DATA ENTRY MODULES ---
        col_input1, col_input2 = st.columns(2)
        
        with col_input1:
            with st.expander("📝 Manual Sentiment Archival", expanded=False):
                with st.form("manual_sentiment_form", clear_on_submit=True, border=False):
                    manual_tag = st.selectbox("Assign to Asset:", tags)
                    manual_date = st.date_input("Review Date:", value=datetime.date.today())
                    f_text = st.text_area("Review Content", placeholder="Paste review text from Google, Yelp, or Surveys...", height=150)
                    
                    if st.form_submit_button("🛡️ AI Score & Vault Entry", use_container_width=True):
                        if f_text:
                            with st.spinner("Analyzing sentiment..."):
                                if archive_sentiment_entry(f_text, manual_tag, manual_date, tenant_id, supabase):
                                    st.success("Entry Scored & Vaulted.")
                                    time.sleep(1)
                                    st.rerun()

        with col_input2:
            with st.expander("📄 Intelligence Bulk Loader (.docx)", expanded=False):
                if not DOCX_AVAILABLE:
                    st.error("System Dependency Missing: Run `pip install python-docx` in your environment to unlock the bulk loader.")
                else:
                    uploaded_doc = st.file_uploader("Upload .docx Source", type="docx")
                    bulk_tag = st.selectbox("Bulk Assign to Asset:", tags)
                    bulk_date = st.date_input("Bulk Review Date:", value=datetime.date.today(), key="bulk_date")
                    
                    if uploaded_doc and st.button("🚀 Execute Bulk AI Parse", use_container_width=True):
                        doc = Document(uploaded_doc)
                        valid_paragraphs = []
                        
                        # Extract Text
                        for para in doc.paragraphs:
                            clean_text = para.text.strip()
                            if len(clean_text) > 20: valid_paragraphs.append(clean_text)
                                
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    clean_text = cell.text.strip()
                                    if len(clean_text) > 20 and clean_text not in valid_paragraphs:
                                        valid_paragraphs.append(clean_text)
                        
                        total_paras = len(valid_paragraphs)
                        
                        if total_paras > 0:
                            success_count, fail_count = 0, 0
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            for idx, text in enumerate(valid_paragraphs):
                                status_text.text(f"Analyzing & Vaulting entry {idx + 1} of {total_paras}...")
                                if archive_sentiment_entry(text, bulk_tag, bulk_date, tenant_id, supabase):
                                    success_count += 1
                                else:
                                    fail_count += 1
                                
                                progress_bar.progress((idx + 1) / total_paras)
                                time.sleep(1.5) # Throttle to respect Gemini API limits
                                
                            status_text.empty()
                            progress_bar.empty()
                            
                            if success_count > 0: st.success(f"✅ Vaulting Complete: {success_count} entries scored.")
                            if fail_count > 0: st.error(f"⚠️ {fail_count} entries failed API analysis.")
                        else:
                            st.warning("No valid review text found in document.")

        st.divider()

        # --- 2. SENTIMENT VAULT RESEARCH FILTERS ---
        f1, f2, f3, f4 = st.columns([1.5, 1, 1.2, 1])
        with f1: search_query = st.text_input("Search Content", placeholder="Keyword search...")
        with f2: filter_asset = st.selectbox("Asset Filter", ["All Assets"] + tags)
        with f3: 
            today = datetime.date.today()
            date_range = st.date_input("Audit Window", value=(today - datetime.timedelta(days=30), today))
        with f4: filter_cat = st.selectbox("Category", ["All", "Positive", "Neutral", "Negative"])

        # --- 3. FETCH AND PROCESS DATA ---
        query = supabase.table("mt_sentiment_history").select("*").eq("tenant_id", tenant_id)
        if filter_asset != "All Assets": query = query.eq("asset", filter_asset)
        if filter_cat != "All": query = query.eq("sentiment_category", filter_cat)
        if isinstance(date_range, tuple) and len(date_range) == 2:
            query = query.gte("timestamp", date_range[0].isoformat()).lte("timestamp", f"{date_range[1].isoformat()}T23:59:59")
        
        res = query.order("timestamp", desc=True).execute()

        if res.data:
            df_vault = pd.DataFrame(res.data)
            if search_query:
                df_vault = df_vault[df_vault['raw_text'].str.contains(search_query, case=False)]

            if not df_vault.empty:
                # --- 4. BENTO METRICS ---
                total_reviews = len(df_vault)
                pos_count = len(df_vault[df_vault['sentiment_category'] == 'Positive'])
                neg_count = len(df_vault[df_vault['sentiment_category'] == 'Negative'])

                st.markdown(f"""
                <div style="display: flex; gap: 1.5rem; margin-bottom: 2rem; margin-top: 1rem;">
                    <div class="bento-card" style="flex: 1; text-align: center;">
                        <p style="color: #6B7280; margin: 0; font-size: 0.85rem; font-weight: 600; text-transform: uppercase;">Total Vault Volume</p>
                        <h2 style="color: #111827; margin: 0.5rem 0; font-size: 2.2rem;">{total_reviews:,}</h2>
                    </div>
                    <div class="bento-card" style="flex: 1; text-align: center; border-bottom: 4px solid #10B981;">
                        <p style="color: #10B981; margin: 0; font-size: 0.85rem; font-weight: 700; text-transform: uppercase;">Positive Nodes</p>
                        <h2 style="color: #111827; margin: 0.5rem 0; font-size: 2.2rem;">{pos_count:,}</h2>
                    </div>
                    <div class="bento-card" style="flex: 1; text-align: center; border-bottom: 4px solid #EF4444;">
                        <p style="color: #EF4444; margin: 0; font-size: 0.85rem; font-weight: 600; text-transform: uppercase;">Critical Nodes</p>
                        <h2 style="color: #111827; margin: 0.5rem 0; font-size: 2.2rem;">{neg_count:,}</h2>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # --- 5. THE FEED ---
                st.markdown("<h4 style='color: #111827;'>Recent Feedback</h4>", unsafe_allow_html=True)
                
                # Limit rendering to 50 for performance
                for _, row in df_vault.head(50).iterrows():
                    cat = row.get('sentiment_category', 'Neutral')
                    score = float(row.get('sentiment_score', 0.0))
                    
                    # Convert -1 to 1 score into a 0 to 100 percentage for the visual bar
                    bar_pct = ((score + 1) / 2) * 100
                    
                    if cat == "Positive": color = "#10B981" # Green
                    elif cat == "Negative": color = "#EF4444" # Red
                    else: color = "#F59E0B" # Gold
                    
                    st.markdown(f"""
                    <div class="bento-card" style="margin-bottom: 1rem; padding: 1.5rem;">
                        <div style="display: flex; justify-content: space-between; align-items: flex-start;">
                            <div style="flex: 4; padding-right: 2rem;">
                                <span style="background-color: #F3F4F6; padding: 0.2rem 0.6rem; border-radius: 4px; font-size: 0.8rem; font-weight: 600; color: #4B5563;">{row.get('asset')}</span>
                                <span style="color: #9CA3AF; font-size: 0.8rem; margin-left: 0.5rem;">{str(row.get('timestamp'))[:10]}</span>
                                <p style="color: #1F2937; margin-top: 0.8rem; line-height: 1.5;">"{row.get('raw_text')}"</p>
                            </div>
                            <div style="flex: 1; text-align: right;">
                                <p style="color: {color}; font-weight: 700; font-size: 1.2rem; margin: 0;">{score:+.2f}</p>
                                <p style="color: #6B7280; font-size: 0.8rem; margin: 0 0 0.5rem 0; text-transform: uppercase;">{cat}</p>
                                <div style="width: 100%; background-color: #E5E7EB; border-radius: 4px; height: 6px;">
                                    <div style="width: {bar_pct}%; background-color: {color}; height: 100%; border-radius: 4px;"></div>
                                </div>
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.info("No records match your search criteria.")
        else:
            st.info("The Sentiment Vault is empty. Upload a document or manually enter a review to begin.")
